#!/usr/bin/env python3
"""
处理 TEM.md 文件，生成紧密格式的 Word 文档
保留词性转换和固定搭配，省略辨析部分
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def process_markdown(file_path):
    """处理 Markdown 文件，提取词性转换和固定搭配"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 匹配词汇条目
    # 格式：### 1. provincial /prəˈvɪnʃl/
    pattern = r'### (\d+)\. (\w+) /([^/]+)/\n\n(.*?)(?=### \d+\. |$)'
    entries = re.findall(pattern, content, re.DOTALL)
    
    processed_entries = []
    
    for entry in entries:
        number, word, phonetic, details = entry
        
        # 提取词性转换
        pos_pattern = r'\*\*1\. 词性转换\*\*\n\n([\s\S]*?)(?=\*\*2\. 固定搭配|$)'
        pos_matches = re.findall(pos_pattern, details)
        pos_content = pos_matches[0].strip() if pos_matches else ""
        
        # 提取固定搭配
        collocation_pattern = r'\*\*2\. 固定搭配/地道表达\*\*\n\n([\s\S]*?)(?=\*\*3\. 重点辨析|$)'
        collocation_matches = re.findall(collocation_pattern, details)
        collocation_content = collocation_matches[0].strip() if collocation_matches else ""
        
        # 处理词性转换内容，移除不必要的换行
        pos_lines = []
        for line in pos_content.split('\n'):
            line = line.strip()
            if line:
                pos_lines.append(line)
        
        # 处理固定搭配内容，移除不必要的换行
        collocation_lines = []
        for line in collocation_content.split('\n'):
            line = line.strip()
            if line:
                collocation_lines.append(line)
        
        processed_entries.append({
            'number': number,
            'word': word,
            'phonetic': phonetic,
            'pos': pos_lines,
            'collocation': collocation_lines
        })
    
    return processed_entries


def create_word_document(entries, output_path):
    """创建 Word 文档"""
    doc = Document()
    
    # 设置页面边距 - 上下边距减小，左右边距保持1厘米
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)  # 减小上边距
        section.bottom_margin = Inches(0.3)  # 减小下边距
        section.left_margin = Inches(0.4)  # 保持左边距1厘米
        section.right_margin = Inches(0.4)  # 保持右边距1厘米
        
        # 设置两栏布局，带分隔线
        from docx.oxml.shared import OxmlElement
        
        # 获取节属性
        sectPr = section._sectPr
        
        # 确保文档只有一个节，这样分栏设置会应用到整个文档
        if len(doc.sections) > 1:
            # 如果有多个节，合并它们
            for i in range(len(doc.sections) - 1, 0, -1):
                doc._body.remove(doc.sections[i]._sectPr)
        
        # 清除现有分栏设置
        for cols in sectPr.xpath('./w:cols'):
            sectPr.remove(cols)
        
        # 添加新的分栏设置
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')  # 两栏
        cols.set(qn('w:space'), '360')  # 栏间距（10pt）
        cols.set(qn('w:sep'), '1')  # 添加分隔线
        cols.set(qn('w:equalWidth'), '1')  # 等宽分栏
        sectPr.append(cols)
    
    # 添加标题
    title = doc.add_heading('TEM-4/8 Vocabs', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 减小标题字号并设置中文字体
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        # 添加中文字体支持
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 设置正文样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(7)  # 更小的字号
    
    # 设置中文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 设置行间距为单倍行距
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(1)  # 段后间距极小
    
    # 遍历处理后的条目，添加到文档
    for entry in entries:
        # 词汇标题：1. provincial /prəˈvɪnʃl/
        heading_text = f"{entry['number']}. {entry['word']} /{entry['phonetic']}/"
        heading = doc.add_paragraph(heading_text, style='Normal')
        
        # 减小标题字号、加粗并设置中文字体
        for run in heading.runs:
            run.font.bold = True
            run.font.size = Pt(8)  # 更小的标题字号
            run.font.name = 'Times New Roman'
            # 添加中文字体支持
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        heading.paragraph_format.line_spacing = 1.0
        heading.paragraph_format.space_after = Pt(0.5)
        
        # 词性转换和固定搭配合并为一段，移除标签文字
        combined_text = ""
        if entry['pos']:
            pos_text = " ".join(entry['pos'])
            pos_text = pos_text.replace('**', '')
            combined_text += f"{pos_text}  "
        
        if entry['collocation']:
            collocation_text = " ".join(entry['collocation'])
            collocation_text = collocation_text.replace('**', '')
            combined_text += f"{collocation_text}"
        
        if combined_text:
            p = doc.add_paragraph(combined_text, style='Normal')
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            
            # 确保中文字体正确显示
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    doc.save(output_path)
    print(f"Word 文档已生成：{output_path}")


if __name__ == "__main__":
    input_file = "d:\\edu-wiki\\docs\\docs\\english\\vocabs\\TEM.md"
    output_file = "d:\\edu-wiki\\TEM_vocab_two_columns.docx"
    
    print("正在处理 Markdown 文件...")
    entries = process_markdown(input_file)
    print(f"共处理 {len(entries)} 个词汇条目")
    
    print("正在生成 Word 文档...")
    create_word_document(entries, output_file)
    print("完成！")
