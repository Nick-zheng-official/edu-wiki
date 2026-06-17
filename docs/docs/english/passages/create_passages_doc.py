from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Read the markdown file
with open('d:/edu-wiki/docs/docs/english/passages/passages.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create new document
doc = Document()

# Set narrow margins for compact layout
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

# Set default font
style = doc.styles['Normal']
font = style.font
font.size = Pt(9)
font.name = 'Times New Roman'

# Parse markdown content
lines = content.split('\n')

def add_run_with_font(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    if bold:
        run.font.bold = True
    return run

def set_columns(section, num_cols, spacing= Cm(0.3)):
    """Set the number of columns for a section"""
    from docx.oxml import OxmlElement
    sectPr = section._sectPr
    # Find or create cols element
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(spacing.twips / 20)))  # w:space is in twentieths of a point

# Set two-column layout after document creation
for section in doc.sections:
    set_columns(section, 2, Cm(0.4))
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

for line in lines:
    line = line.strip()
    
    if not line:
        # Empty line - add small spacing
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        continue
    
    # Headers
    if line.startswith('# '):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = add_run_with_font(p, line[2:], bold=True)
        run.font.size = Pt(14)
    elif line.startswith('## '):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = add_run_with_font(p, line[3:], bold=True)
        run.font.size = Pt(12)
    elif line.startswith('### '):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = add_run_with_font(p, line[4:], bold=True)
        run.font.size = Pt(10)
    elif line.startswith('---'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.add_run('─' * 35)
    elif line.startswith('**') and line.endswith('**'):
        # Bold line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = add_run_with_font(p, line[2:-2], bold=True)
        run.font.size = Pt(10)
    elif line.startswith('_') and line.endswith('_'):
        # Italic line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line[1:-1])
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.italic = True
    elif line.startswith('|'):
        # Table row - skip for now
        continue
    elif line.startswith('>'):
        # Quote
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        run = add_run_with_font(p, line[1:])
    elif line.startswith('```'):
        # Code block - skip
        continue
    else:
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 0.9
        
        # Process inline formatting
        i = 0
        while i < len(line):
            if line[i:i+2] == '**':
                # Bold
                end = line.find('**', i+2)
                if end != -1:
                    run = p.add_run(line[i+2:end])
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    i = end + 2
                else:
                    run = p.add_run(line[i])
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    i += 1
            elif line[i] == '*':
                # Italic
                end = line.find('*', i+1)
                if end != -1:
                    run = p.add_run(line[i+1:end])
                    run.font.italic = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    i = end + 1
                else:
                    run = p.add_run(line[i])
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    i += 1
            else:
                # Regular text
                run = p.add_run(line[i])
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                i += 1

# Add page numbers to footer
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add "Page X" text
run = footer_para.add_run('Page ')
run.font.size = Pt(9)
run.font.name = 'Times New Roman'

# Add page number field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')

fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

run2 = footer_para.add_run()
run2.font.size = Pt(9)
run2.font.name = 'Times New Roman'
run2._r.append(fldChar1)
run2._r.append(instrText)
run2._r.append(fldChar2)
run2._r.append(fldChar3)

# Save
output_path = 'd:/edu-wiki/docs/docs/english/passages/passages_compact.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
